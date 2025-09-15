# resume_service.py
import os
import io
import uuid
import re
from typing import List, Tuple, Optional, Dict
import json
import re

from flask import Blueprint, request, jsonify, send_file
from dotenv import load_dotenv

# DB
from supabase import create_client, Client

from copy import deepcopy
import string

# DOCX
from docx import Document
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor
from docx.enum.style import WD_STYLE_TYPE

# Similarity
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from transformers import pipeline
import google.generativeai as genai

# Optional embeddings (better ranking if available)
try:
    from sentence_transformers import SentenceTransformer
    _EMBED = SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")
except Exception:
    _EMBED = None

import requests

load_dotenv()
resume_bp = Blueprint("resume_bp", __name__)

SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_SERVICE_ROLE_KEY") or os.getenv("SUPABASE_ANON_KEY")
supabase: Optional[Client] = create_client(SUPABASE_URL, SUPABASE_KEY) if (SUPABASE_URL and SUPABASE_KEY) else None

OLLAMA_BASE = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "llama3.1:8b")

DOWNLOAD_DIR = os.path.join(os.getcwd(), "generated_resumes")
os.makedirs(DOWNLOAD_DIR, exist_ok=True)

# Load the Hugging Face model for paraphrasing
_paraphraser = pipeline("text2text-generation", model="t5-small", device=0)  # Use device=0 for GPU, -1 for CPU

# Configure Gemini API
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
model = genai.GenerativeModel('gemini-1.5-flash-latest')

# ----------------------------
# Utilities
# ----------------------------
BULLET_STYLES = {"List Paragraph", "List Bullet", "List Bullet 2", "List Bullet 3", "List Continue"}

# Replace this
# HEADING_NAMES = {"heading 1", "heading 2", "heading 3"}
# def _is_heading(...):

HEADING_NAMES = {"heading 1"}  # only H1 is a section heading

def _is_heading(p: Paragraph) -> bool:
    """Top-level section heading: Heading 1 OR ALL-CAPS short text (e.g., WORK EXPERIENCE)."""
    try:
        if p.style and p.style.name and p.style.name.lower() in HEADING_NAMES:
            return True
    except Exception:
        pass
    txt = (p.text or "").strip()
    # ALL-CAPS, short, not just punctuation
    if (
        3 <= len(txt) <= 40 and
        txt.upper() == txt and
        re.sub(r"[^A-Za-z ]", "", txt).strip()  # has letters
    ):
        return True
    return False

def _find_section(doc: Document, label: str) -> tuple[int, int]:
    """(start_idx, end_idx_exclusive) for the top-level section whose heading contains label."""
    label_low = label.lower()
    heads = [i for i, p in enumerate(doc.paragraphs) if _is_heading(p)]
    heads.append(len(doc.paragraphs))  # sentinel
    for idx in range(len(heads) - 1):
        h = heads[idx]
        txt = (doc.paragraphs[h].text or "").strip().lower()
        if label_low in txt:
            return h, heads[idx + 1]
    return -1, -1


def _insert_para_after(paragraph: Paragraph, text: str, style_name: str | None = None) -> Paragraph:
    from docx.oxml import OxmlElement
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style_name:
        try:
            new_para.style = style_name
        except Exception:
            pass
    run = new_para.add_run(text)
    return new_para

def _cap_words(s: str, max_words: int = 20) -> str:
    w = (s or "").split()
    return " ".join(w[:max_words]) + ("…" if len(w) > max_words else "")

def _add_bullet_paragraph(after_para, text, style_name: str | None = None):
    p = _insert_para_after(after_para, "", style_name or None)
    run = p.add_run("• " + _clean(text))  # ⟵ no _cap_words
    try:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    except Exception:
        pass
    try:
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
    except Exception:
        pass
    return p


def _is_bullet(p: Paragraph) -> bool:
    """Detect bullets by style, numbering, or literal bullet/dash prefix."""
    try:
        if p.style and p.style.name in BULLET_STYLES:
            return True
        # numbering/bullets via XML
        if p._p is not None and getattr(p._p, "pPr", None) is not None and getattr(p._p.pPr, "numPr", None) is not None:
            return True
    except Exception:
        pass
    # literal bullets/dashes (common after PDF conversions)
    txt = (p.text or "").lstrip()
    if txt.startswith(("•", "-", "–", "—")) and len(txt) > 2:
        return True
    return False

def _clean(s: str) -> str:
    s = re.sub(r"\s+", " ", (s or "").strip())
    # trim common suffix artifacts like "— on-site"
    s = re.sub(r"\s*[—-]\s*on-?site\s*$", "", s, flags=re.I)
    return s

def _shorten_to_words(s: str, max_words: int = 24) -> str:
    words = (s or "").split()
    return " ".join(words[:max_words]) + ("…" if len(words) > max_words else "")

def _delete_paragraph(p: Paragraph):
    p._element.getparent().remove(p._element)
    p._p = p._element = None

def _get_all_paragraph_text(doc: Document) -> List[str]:
    return [p.text for p in doc.paragraphs]

def _extract_bullets(doc: Document) -> List[Tuple[int, Paragraph]]:
    bullets: List[Tuple[int, Paragraph]] = []
    for idx, p in enumerate(doc.paragraphs):
        if _is_bullet(p):
            bullets.append((idx, p))
    return bullets

def _slug(s: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", (s or "").lower())

def _append_dates_inline(p: Paragraph, dates: str):
    if not dates:
        return
    r = p.add_run(" — " + dates)   # em dash-ish
    r.bold = False
    r.italic = False
    r.font.name = "Garamond"
    r.font.size = Pt(10)

def _experience_dates_map(raw_list: list[str]) -> dict[str, str]:
    """Map normalized experience title -> 'Mon YYYY – Mon YYYY|Ongoing|Present'"""
    out = {}
    for block in raw_list or []:
        first = (block.splitlines() or [""])[0]
        line = " ".join(first.split())  # collapse tabs/spaces
        m = re.search(
            r"(.*?)(?:\s{2,}|\t+)\s*([A-Za-z]{3,9}\s+\d{4}\s*[–-]\s*(?:[A-Za-z]{3,9}\s+\d{4}|Ongoing|Present))$",
            line
        )
        if m:
            title, dates = m.group(1).strip(" ,"), m.group(2)
        else:
            m2 = re.search(r"([A-Za-z]{3,9}\s+\d{4}\s*[–-]\s*(?:[A-Za-z]{3,9}\s+\d{4}|Ongoing|Present))", first)
            dates = m2.group(1) if m2 else ""
            title = re.sub(r"\s+", " ", first.replace(dates, "")).strip(" ,")
        slug = _slug(title)
        if slug:
            out[slug] = dates
    return out


def _style_title(p: Paragraph, *, blue: bool = False):
    # Make all runs Garamond 10, bold; optionally blue
    if not p.runs:
        p.add_run("")  # ensure there is at least one run
    for r in p.runs:
        r.font.name = "Garamond"
        r.font.size = Pt(10)
        r.bold = True
        if blue:
            r.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)  # deep-ish blue

def _append_skills_inline(p: Paragraph, skills: list[str]):
    if not skills: 
        return
    txt = " | " + ", ".join(skills)
    r = p.add_run(txt)
    r.italic = True
    r.bold = False
    r.font.name = "Garamond"
    r.font.size = Pt(10)


def _para_index(doc: Document, target: Paragraph) -> int:
    """Find a paragraph index by comparing the underlying XML element."""
    for i, p in enumerate(doc.paragraphs):
        if p._p is target._p:
            return i
    return -1

def _embed(texts: List[str]) -> np.ndarray:
    """Return embeddings (MiniLM if available; TF-IDF fallback is only used when both sides share the same vectorizer).
    IMPORTANT: Do not use TF-IDF embeddings for cross-set cosine unless fitted jointly."""
    if _EMBED:
        return np.array(_EMBED.encode(texts, normalize_embeddings=True))
    # If no SentenceTransformer, caller should avoid using this channel across different corpora.
    # We still return a TF-IDF embedding for same-corpus cosine, but cross-corpus is handled elsewhere.
    vec = TfidfVectorizer().fit(texts)
    mat = vec.transform(texts).astype(float).toarray()
    # L2 normalize
    mat = mat / (np.linalg.norm(mat, axis=1, keepdims=True) + 1e-9)
    return mat

def _cos_sim(a: np.ndarray, b: np.ndarray) -> np.ndarray:
    return a @ b.T

def _extract_keywords(text: str, topk: int = 30) -> List[str]:
    vec = TfidfVectorizer(ngram_range=(1, 2), stop_words="english", max_features=2000)
    X = vec.fit_transform([text or ""])
    row = X.toarray()[0]
    terms = np.array(vec.get_feature_names_out())
    inds = np.argsort(-row)
    kws = [terms[i] for i in inds[:topk] if row[i] > 0]
    kws = [k for k in kws if 2 <= len(k) <= 30][:topk]
    return kws

def _rewrite_with_ollama(bullet: str, jd_text: str, allowed_terms: List[str]) -> str:
    system = (
        "You are rewriting resume bullets to better match a job description.\n"
        "Rules:\n"
        "- Focus on aligning the bullet with the job description.\n"
        "- Highlight relevant skills, tools, and achievements that match the job description.\n"
        "- Do NOT invent experience or tools the candidate never used.\n"
        "- Keep one sentence, punchy, <= 24 words, action-oriented.\n"
        "- Keep the original meaning; just align phrasing to JD terminology.\n"
        "- No emojis, no fluff."
    )
    prompt = (
        f"Job description (excerpt):\n{jd_text[:1200]}\n\n"
        f"Allowed terms (safe to mention): {', '.join(sorted(set(allowed_terms)))}\n\n"
        f"Original bullet:\n\"{bullet}\"\n\n"
        "Rewrite now:"
    )
    try:
        r = requests.post(
            f"{OLLAMA_BASE}/api/generate",
            json={"model": OLLAMA_MODEL, "prompt": f"System:\n{system}\n\nUser:\n{prompt}\n"},
            timeout=60,
        )
        r.raise_for_status()
        out = (r.json() or {}).get("response", "").strip()
        out = out.replace("\n", " ")
        print(f"Sending request to Ollama: {prompt}")
        print(f"Ollama response: {out}")
        return _shorten_to_words(_clean(out), 24) or bullet
    except Exception:
        # fallback: just return a shortened original
        return _shorten_to_words(bullet, 24)

def _rewrite_with_huggingface(bullet: str, jd_text: str) -> str:
    """
    Rewrite a resume bullet to better align with the job description using a Hugging Face model.
    """
    try:
        # Create a prompt for the model
        prompt = (
            f"Paraphrase the following resume bullet to better align with the job description:\n\n"
            f"Job Description: {jd_text[:500]}\n\n"
            f"Resume Bullet: {bullet}\n\n"
            f"Rewritten Bullet:"
        )
        # Generate the paraphrased bullet
        response = _paraphraser(prompt, max_length=50, num_return_sequences=1)
        rewritten_bullet = response[0]["generated_text"]
        return _shorten_to_words(_clean(rewritten_bullet), 24)
    except Exception as e:
        print(f"⚠️ Hugging Face rewrite failed: {e}")
        return bullet  # Fallback to the original bullet

def _rewrite_no_llm(bullet: str, jd_text: str, resume_terms: List[str]) -> str:
    """
    Rewrite a resume bullet to better align with the job description using cosine similarity.
    """
    # Extract sentences from the job description
    jd_sents = _sentences(jd_text)
    if not jd_sents:
        return bullet  # No sentences to compare with

    # Compute TF-IDF embeddings for the bullet and JD sentences
    corpus = [bullet] + jd_sents
    vec = TfidfVectorizer(ngram_range=(1, 2), stop_words="english", min_df=1)
    X = vec.fit_transform(corpus)  # TF-IDF matrix
    bullet_vec = X[0]  # First row is the bullet
    jd_vecs = X[1:]    # Remaining rows are JD sentences

    # Compute cosine similarity between the bullet and each JD sentence
    sims = (bullet_vec @ jd_vecs.T).toarray().flatten()
    if not sims.size:
        return bullet  # No similarity scores

    # Find the most similar JD sentence
    best_idx = sims.argmax()
    best_sent = jd_sents[best_idx]
    best_score = sims[best_idx]

    # If similarity is high, reword the bullet to align with the best JD sentence
    if best_score > 0.3:  # Threshold for similarity
        candidate = f"{bullet} (aligned with: {best_sent})"
    else:
        # If similarity is low, just append relevant keywords
        jd_kws = set(_extract_keywords(jd_text, topk=25))
        res_kws = set(_extract_keywords(" ".join(resume_terms), topk=25))
        overlap = [k for k in jd_kws if k in res_kws and len(k) > 2]
        add = [k for k in overlap if k.lower() not in bullet.lower()][:2]
        candidate = re.sub(r"[.;:,\s]+$", "", bullet) + f" using {', '.join(add)}." if add else bullet

    return _shorten_to_words(_clean(candidate), 24)

def _collect_resume_terms(doc: Document) -> List[str]:
    # super-light skills collector: grab tokens across resume
    text = " ".join(_get_all_paragraph_text(doc))
    terms = re.findall(r"[A-Za-z][A-Za-z0-9+.#-]{1,32}", text)
    return terms[:300]

def _fit_to_one_page(doc: Document, keep_order: bool = True, max_bullets: int = 18, top_bullets: int = 3):
    """Ensure top bullets are always included and cap total bullet count."""
    bullets = _extract_bullets(doc)
    if len(bullets) <= max_bullets:
        return

    # Sort bullets by relevance
    bullets_sorted = sorted(bullets, key=lambda t: getattr(t[1], "_relevance", 0.0), reverse=True)

    # Ensure top bullets are always included
    top_bullets_set = set(idx for idx, _ in bullets_sorted[:top_bullets])
    to_keep = top_bullets_set.union(idx for idx, _ in bullets_sorted[:max_bullets])

    # Remove bullets not in the keep set
    for idx, p in bullets:
        if idx not in to_keep:
            _delete_paragraph(p)

def _sentences(text: str) -> List[str]:
    # Coarse sentence split: terminals, newlines, semicolons
    raw = re.split(r'(?<=[.!?])\s+|\n{1,}|\s*;\s*', text or "")
    sents = [_clean(s) for s in raw if 5 <= len(s.split()) <= 60]
    if not sents:
        sents = [_clean(s) for s in raw if _clean(s)]
    return sents[:120]  # cap for speed

def _scale_01(arr: np.ndarray) -> np.ndarray:
    if arr.size == 0:
        return arr
    lo = float(np.percentile(arr, 10))
    hi = float(np.percentile(arr, 90))
    if hi <= lo + 1e-8:
        return np.clip((arr - lo), 0, 1)  # all same → zeros
    return np.clip((arr - lo) / (hi - lo), 0, 1)

def _semantic_maxsim_per_bullet(bullets: List[str], jd_text: str) -> np.ndarray:
    """Max cosine similarity per bullet vs JD sentences using SentenceTransformer if available.
    If ST is not available, return zeros and let the lexical channel handle it."""
    jd_sents = _sentences(jd_text)
    if not jd_sents or not bullets:
        return np.zeros(len(bullets), dtype=float)
    if _EMBED is None:
        # Without a shared embedding space, skip semantic channel
        return np.zeros(len(bullets), dtype=float)
    em_bul = _embed(bullets)
    em_jd = _embed(jd_sents)
    sims = _cos_sim(em_bul, em_jd)
    return sims.max(axis=1) if sims.size else np.zeros(len(bullets), dtype=float)

def _lexical_maxsim_per_bullet(bullets: List[str], jd_text: str) -> np.ndarray:
    """TF-IDF (1,2-gram) cosine vs JD sentences; take max per bullet."""
    jd_sents = _sentences(jd_text)
    if not jd_sents or not bullets:
        return np.zeros(len(bullets), dtype=float)
    corpus = bullets + jd_sents
    vec = TfidfVectorizer(ngram_range=(1, 2), stop_words="english", min_df=1)
    X = vec.fit_transform(corpus)  # l2-normalized rows by default
    B = X[:len(bullets), :]
    J = X[len(bullets):, :]
    sims = (B @ J.T).toarray()  # cosine for l2-normalized tfidf
    return sims.max(axis=1) if sims.size else np.zeros(len(bullets), dtype=float)

def _hybrid_scores(bullets: List[str], jd_text: str, w_sem: float = 0.7) -> Tuple[np.ndarray, Dict]:
    """Blend semantic and lexical max-sim (both scaled 0–1)."""
    sem = _semantic_maxsim_per_bullet(bullets, jd_text)  # 0..1 (zeros if no ST)
    lex = _lexical_maxsim_per_bullet(bullets, jd_text)   # ~0..1
    sem_n = _scale_01(sem)
    lex_n = _scale_01(lex)
    combo = w_sem * sem_n + (1.0 - w_sem) * lex_n
    meta = {
        "semantic_backend": "sentence-transformers" if _EMBED else "lexical-only",
        "semantic_raw_min": float(sem.min()) if sem.size else 0.0,
        "semantic_raw_max": float(sem.max()) if sem.size else 0.0,
    }
    return combo, meta

# ----------------------------
# Core route
# ----------------------------
@resume_bp.route("/api/resume/tune", methods=["POST"])
def tune_resume():
    """
    Generate a tailored resume using Supabase data and the job description.
    """
    try:
        # Check if the file is uploaded
        if "file" not in request.files:
            return jsonify({"success": False, "error": "No resume file uploaded"}), 400

        file = request.files["file"]
        if not file.filename.lower().endswith(".docx"):
            return jsonify({"success": False, "error": "Please upload a .docx file"}), 400

        # Get user_id and job_description from the request
        user_id = request.form.get("user_id", "")
        job_description = request.form.get("job_description", "")

        if not user_id or not job_description:
            return jsonify({"success": False, "error": "Missing user_id or job_description"}), 400

        # Fetch user data from Supabase
        if not supabase:
            return jsonify({"success": False, "error": "Database not available"}), 500

        user_query = supabase.table("Users").select("*").eq("user_uuid", user_id).limit(1).execute()
        if not user_query.data:
            return jsonify({"success": False, "error": "User not found"}), 404

        user_info = user_query.data[0]
        projects = user_info.get("Projects", [])
        experiences = user_info.get("Experiences", [])
        exp_dates = _experience_dates_map(experiences)

        # Debug: Log the fetched projects and experiences
        print(f"📚 Supabase Projects: {projects}")
        print(f"📚 Supabase Experiences: {experiences}")

        # Save the uploaded resume
        resume_bytes = io.BytesIO(file.read())
        doc = Document(resume_bytes)

        # Generate tailored content using Gemini
        tailored_projects = _generate_tailored_points(projects, job_description, "project")
        tailored_experiences = _generate_tailored_points(experiences, job_description, "experience")

        print("✅ Tailored Projects:", tailored_projects)
        print("✅ Tailored Experiences:", tailored_experiences)

        # Add tailored content to the resume
        tailored_path, summary = process_resume(doc, tailored_projects, tailored_experiences, exp_dates)

        # Save the tailored resume to a file
        out_id = str(uuid.uuid4())[:8]
        out_name = f"tailored-{out_id}.docx"
        out_path = os.path.join(DOWNLOAD_DIR, out_name)
        tailored_path.save(out_path)

        return jsonify({
            "success": True,
            "download_url": f"/api/resume/download/{out_name}",
            "change_summary": summary
        })

    except Exception as e:
        print(f"❌ Error in tune_resume: {e}")
        return jsonify({"success": False, "error": str(e)}), 500

@resume_bp.route("/api/resume/download/<fname>", methods=["GET"])
def download(fname: str):
    path = os.path.join(DOWNLOAD_DIR, fname)
    if not os.path.exists(path):
        return jsonify({"success": False, "error": "Not found"}), 404
    return send_file(
        path,
        as_attachment=True,
        download_name=fname,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

def process_resume(doc: Document, projects: list[dict], experiences: list[dict], exp_dates: dict[str, str] | None = None):
    """
    Replace 'WORK EXPERIENCE' and 'PROJECTS/TECHNICAL PROJECTS' bodies with tailored items.
    Header/Education/Skills untouched. Creates Projects section if missing.
    """
    summary: list[str] = []
    exp_dates = exp_dates or {}

    def _insert_entry(after_para: Paragraph, name: str, bullets: list[str]) -> Paragraph:
        # Title as bold Normal, not Heading 2 (so it won't be treated as a section heading)
        title_p = _insert_para_after(after_para, _clean(name), None)
        for r in title_p.runs:
            r.bold = True
        after = title_p
        for b in bullets:
            after = _add_bullet_paragraph(after, b, None)  # manual • with hanging indent
        return after

    def _clear_section_body(s_idx: int, e_idx: int) -> None:
        # delete everything between heading and next top-level heading
        for i in range(e_idx - 1, s_idx, -1):
            _delete_paragraph(doc.paragraphs[i])

    def _replace_section(label: str, items: list[dict], *, is_projects: bool = False) -> int:
        s, e = _find_section(doc, label)
        if s == -1:
            return 0

        # clear body
        for idx in range(e - 1, s, -1):
            _delete_paragraph(doc.paragraphs[idx])

        after = doc.paragraphs[s]
        written = 0
        for item in items:
            name = _clean(item.get("name", ""))
            bullets = [b for b in (item.get("bullets") or []) if _clean(b)]
            skills = item.get("skills") or []
            if not name or not bullets:
                continue

            # entry title
            title_p = _insert_para_after(after, name, None)
            _style_title(title_p, blue=False)
            if not is_projects:
                _append_dates_inline(title_p, exp_dates.get(_slug(name), ""))
            else:
                if skills:
                    _append_skills_inline(title_p, skills)
            after = title_p


            # bullets (keep your existing bullet renderer)
            for b in bullets:
                after = _add_bullet_paragraph(after, b, None)

            written += 1

        print(f"✅ Wrote {written} items into section '{label}'.")
        return written

    def _ensure_projects_section(heading_text: str = "TECHNICAL PROJECTS") -> int:
        ps, pe = _find_section(doc, "project")
        if ps != -1:
            # restyle existing heading
            head_p = doc.paragraphs[ps]
            _style_title(head_p, blue=True)
            return ps

        # insert after Experience if present, else at end
        es, ee = _find_section(doc, "experience")
        after_para = doc.paragraphs[ee - 1] if es != -1 else doc.paragraphs[-1]
        new_head = _insert_para_after(after_para, heading_text.upper(), None)
        _style_title(new_head, blue=True)
        return _para_index(doc, new_head)
    
        # --- EXPERIENCE first ---
    exp_written  = _replace_section("experience", experiences, is_projects=False)

    # --- PROJECTS next (use existing 'PROJECTS' or 'TECHNICAL PROJECTS'; create if missing) ---
    proj_written = (_replace_section("projects", projects, is_projects=True) or _replace_section("project",  projects, is_projects=True))
    if proj_written == 0 and projects:
        _ensure_projects_section("TECHNICAL PROJECTS")
        proj_written = (_replace_section("project", projects, is_projects=True) or
                        _replace_section("projects", projects, is_projects=True))
    if exp_written:
        summary.append(f"Replaced Experience with {exp_written} item(s)")
    else:
        summary.append("Experience section not found; left unchanged")
    if proj_written:
        summary.append(f"Replaced/created Projects with {proj_written} item(s)")
    else:
        summary.append("Projects section not found; left unchanged")

    return doc, summary


def _extract_json_block(s: str) -> str:
    """Pull the first JSON array/object from a text blob (handles code fences)."""
    # Remove common code fences
    s = re.sub(r"^```(?:json)?|```$", "", s.strip(), flags=re.I|re.M)
    # Greedy find of the first [... ] or {...}
    m = re.search(r"(\{.*\}|\[.*\])", s, flags=re.S)
    return m.group(1) if m else ""

def _generate_tailored_points(items, job_description, item_type):
    if not items:
        print(f"No {item_type}s provided for tailoring.")
        return []

    prompt = f"""
You are a senior software-engineering resume writer optimizing for ATS. Select the most professional, outstanding, and impactful {item_type}s.

Return a STRICT JSON array (no prose, no markdown). Each element:
{{
  "name": "<{item_type} name>",
  "skills": ["<tech/skill 1>", "..."],             // 5–10 items; only real tech from the item/JD
  "bullets": ["<b1>", "<b2>", "<b3>", "<b4>"]      // 3–4 bullets allowed (send 3–4)
}}

Bullet rules (VERY IMPORTANT):
- Action-first verbs (Designed, Implemented, Optimized, Architected, Shipped, Automated, Refactored, Scaled).
- 18–40 words; clear problem → approach → impact with credible metrics when implied (%, pXX latency, QPS, users, cost, memory/CPU).
- Use precise SWE terms: REST/GraphQL endpoints, data modeling, indexing, transactions/ACID, caching (Redis), queues/streams, batch vs. streaming, concurrency, profiling, CI/CD, containerization, IaC, monitoring/alerting.
- Name technologies naturally (frameworks, DBs, cloud, testing, pipelines). No emojis. No fabrication.

Job Description:
{job_description}

Candidate {item_type.title()}s (raw list):
{items}

Output JSON ONLY.
""".strip()

    try:
        response = model.generate_content(prompt)
        raw = (response.text or "").strip()
        js = _extract_json_block(raw)
        data = json.loads(js)

        cleaned = []
        for obj in data[:3]:
            name = (obj.get("name") or "").strip()
            skills = [s.strip() for s in (obj.get("skills") or []) if s.strip()][:10]
            bullets = [b.strip().strip("-•–— ").strip() for b in (obj.get("bullets") or [])]
            bullets = [b for b in bullets if b][:4]    # ⟵ up to 4
            if name and len(bullets) >= 2:             # ⟵ at least 2
                cleaned.append({"name": name, "skills": skills, "bullets": bullets})

        print(f"✅ Parsed Tailored {item_type.title()}s (JSON): {cleaned}")
        return cleaned
    except Exception as e:
        print(f"⚠️ JSON parse failed for {item_type}s: {e}\nRaw:\n{response.text if 'response' in locals() else ''}")
        return []


def _parse_listy_text(output: str, item_type: str):
    lines = output.splitlines()
    tailored_points = []
    current = None

    for raw in lines:
        if not raw.strip():
            continue
        indent = len(raw) - len(raw.lstrip(" "))
        text = raw.strip()

        # Headings that look like: "- ProjectName" at indent 0
        if indent == 0 and re.match(r"^- +\S", text):
            if current:
                tailored_points.append(current)
            current = {"name": text[2:].strip(), "bullets": []}
            continue

        # Markdown bold headings: "**Name**" or "* Name"
        if re.match(r"^\*+\s*\S", text):
            name = re.sub(r"^\*+|\*+$", "", text).strip()
            if current:
                tailored_points.append(current)
            current = {"name": name, "bullets": []}
            continue

        # Numbered headings: "1. Name"
        if re.match(r"^\d+\.\s+\S", text) and indent == 0:
            if current:
                tailored_points.append(current)
            current = {"name": re.sub(r"^\d+\.\s+", "", text).strip(), "bullets": []}
            continue

        # Bullets: indented dash/•/–/—, or any dash when we already have a current item
        if (indent >= 2 and re.match(r"^[-•–—]\s+\S", text)) or (current and re.match(r"^[-•–—]\s+\S", text)):
            bullet = re.sub(r"^[-•–—]\s+", "", text).strip()
            if current:
                current["bullets"].append(bullet)
            continue

        # If it's a non-bullet line and we already have an item, treat as a heading line
        if current is None and text:
            current = {"name": text, "bullets": []}

    if current:
        tailored_points.append(current)

    # Clean & clamp
    cleaned = []
    for obj in tailored_points:
        name = (obj.get("name") or "").strip()
        bullets = [b.strip().strip("-•–— ").strip() for b in obj.get("bullets", []) if b.strip()]
        bullets = bullets[:3]
        if name and bullets:
            cleaned.append({"name": name, "bullets": bullets})

    print(f"✅ Parsed Tailored {item_type.title()}s (fallback): {cleaned}")
    return cleaned[:3]


def _create_list_bullet_style(doc):
    """
    Create a "List Bullet" style in the document if it doesn't exist.
    """
    try:
        styles = doc.styles
        if "List Bullet" not in styles:
            bullet_style = styles.add_style("List Bullet", WD_STYLE_TYPE.PARAGRAPH)
            bullet_style.font.name = "Calibri"
            bullet_style.font.size = Pt(11)
            bullet_style.paragraph_format.left_indent = Pt(18)
            bullet_style.paragraph_format.space_before = Pt(6)
            bullet_style.paragraph_format.space_after = Pt(6)
            bullet_style.paragraph_format.first_line_indent = Pt(-18)
            print("✅ Created 'List Bullet' style.")
    except Exception as e:
        print(f"⚠️ Error creating 'List Bullet' style: {e}")
